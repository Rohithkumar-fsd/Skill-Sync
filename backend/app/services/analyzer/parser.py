"""
In-memory resume parser.

Supports PDF (PyMuPDF), DOCX (python-docx), and plain TXT.
No files are written to disk – all processing happens on the UploadFile byte stream.
"""

import io
from fastapi import UploadFile, HTTPException


SUPPORTED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Reads the UploadFile entirely into RAM and extracts raw text.
    Raises HTTP 415 for unsupported formats.
    """
    raw_bytes = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return _parse_pdf(raw_bytes)
    elif filename.endswith(".docx"):
        return _parse_docx(raw_bytes)
    elif filename.endswith(".txt"):
        return raw_bytes.decode("utf-8", errors="ignore").strip()
    else:
        raise HTTPException(
            status_code=415,
            detail=(
                "Unsupported resume format. "
                "Please upload a .pdf, .docx, or .txt file."
            ),
        )


# ─── Private helpers ──────────────────────────────────────────────────────────

def _parse_pdf(data: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="PyMuPDF not installed. Run: pip install PyMuPDF",
        ) from exc

    doc = fitz.open(stream=data, filetype="pdf")
    pages = [page.get_text("text") for page in doc]
    return "\n".join(pages).strip()


def _parse_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Run: pip install python-docx",
        ) from exc

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    # Paragraphs (main body + headings)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables (skills tables are common in resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n".join(parts).strip()
